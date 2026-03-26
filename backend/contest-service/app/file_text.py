import io
import zipfile

CODE_EXTENSIONS = {
    ".py",
    ".js",
    ".ts",
    ".jsx",
    ".tsx",
    ".java",
    ".c",
    ".cpp",
    ".h",
    ".cs",
    ".go",
    ".rs",
    ".rb",
    ".php",
    ".swift",
    ".kt",
    ".scala",
    ".html",
    ".css",
    ".scss",
    ".json",
    ".yaml",
    ".yml",
    ".toml",
    ".md",
    ".txt",
    ".sh",
    ".sql",
}
_MAX_ZIP_FILE_BYTES = 50_000
_MAX_ZIP_TOTAL_BYTES = 200_000
_IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}


def extract_text_pdf(data: bytes) -> str:
    import fitz

    doc = fitz.open(stream=data, filetype="pdf")
    return "\n".join(page.get_text() for page in doc).strip()


def extract_text_docx(data: bytes) -> str:
    from docx import Document

    doc = Document(io.BytesIO(data))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
    return "\n".join(parts)


def extract_text_zip(data: bytes) -> str:
    parts = []
    total = 0
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        for info in zf.infolist():
            if info.is_dir():
                continue
            name = info.filename.lower()
            ext = ("." + name.rsplit(".", 1)[-1]) if "." in name else ""
            file_data = zf.read(info.filename)
            text: str | None = None
            if ext in CODE_EXTENSIONS:
                try:
                    text = file_data[:_MAX_ZIP_FILE_BYTES].decode("utf-8", errors="replace")
                except Exception:
                    continue
            elif ext == ".pdf":
                try:
                    text = extract_text_pdf(file_data)
                except Exception:
                    continue
            elif ext == ".docx":
                try:
                    text = extract_text_docx(file_data)
                except Exception:
                    continue
            if text:
                parts.append(f"=== {info.filename} ===\n{text}")
                total += len(text)
                if total >= _MAX_ZIP_TOTAL_BYTES:
                    parts.append("... (truncated)")
                    break
    return "\n\n".join(parts)


_MAX_ZIP_IMAGES = 10


def extract_zip_images(data: bytes) -> list[tuple[str, bytes]]:
    """Возвращает список (имя_файла, байты) для изображений внутри ZIP (макс. 10)."""
    result = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            for info in zf.infolist():
                if info.is_dir():
                    continue
                name = info.filename.lower()
                ext = ("." + name.rsplit(".", 1)[-1]) if "." in name else ""
                if ext in _IMAGE_EXTENSIONS:
                    result.append((info.filename, zf.read(info.filename)))
                    if len(result) >= _MAX_ZIP_IMAGES:
                        break
    except Exception:
        pass
    return result


def extract_file_text(filename: str, data: bytes) -> str | None:
    name = filename.lower()
    try:
        if name.endswith(".pdf"):
            return extract_text_pdf(data)
        elif name.endswith(".docx"):
            return extract_text_docx(data)
        elif name.endswith(".zip"):
            return extract_text_zip(data)
    except Exception:
        pass
    return None
